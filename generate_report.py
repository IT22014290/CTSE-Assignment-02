"""
Generate SE4010 CTSE Assignment 2 Word Report
Run: python3 generate_report.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy
from pathlib import Path

doc = Document()

# ── Page margins ────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)


# ── Helpers ──────────────────────────────────────────────────────────────────

def set_font(run, name="Calibri", size=11, bold=False, italic=False,
             color=None):
    run.font.name  = name
    run.font.size  = Pt(size)
    run.bold       = bold
    run.italic     = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def heading(text, level=1, color=(31, 73, 125)):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = RGBColor(*color)
        run.font.name = "Calibri"
    return p


def para(text="", bold=False, italic=False, size=11, color=None,
         align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    if text:
        run = p.add_run(text)
        set_font(run, bold=bold, italic=italic, size=size, color=color)
    return p


def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    set_font(run)
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    return p


def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.name = "Calibri"
            run.font.size = Pt(10)
        # Blue background
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "1F497D")
        tc_pr.append(shd)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = str(cell_text)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in cell.paragraphs[0].runs:
                run.font.name = "Calibri"
                run.font.size = Pt(10)
        # Alternating row shading
        if r_idx % 2 == 0:
            for cell in row.cells:
                tc_pr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "DCE6F1")
                tc_pr.append(shd)

    # Column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()
    return table


def separator():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1F497D")
    pBdr.append(bottom)
    pPr.append(pBdr)


def code_block(text):
    """Light-grey shaded monospace paragraph for code / CLI commands."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    # light grey shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    return p


def screenshot_placeholder(label, hint=""):
    """Adds a grey box with a label as a placeholder for a screenshot."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "D9D9D9")
    pPr.append(shd)
    run = p.add_run(f"[ SCREENSHOT: {label} ]")
    run.bold = True
    run.font.color.rgb = RGBColor(64, 64, 64)
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    if hint:
        h = doc.add_paragraph()
        hr = h.add_run(f"How to get: {hint}")
        hr.italic = True
        hr.font.size = Pt(9)
        hr.font.color.rgb = RGBColor(100, 100, 100)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()


# ============================================================================
# COVER PAGE
# ============================================================================

doc.add_paragraph()
doc.add_paragraph()

p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p_title.add_run("SE4010 – Cloud Technologies and Software Engineering")
set_font(run, size=14, bold=True, color=(31, 73, 125))

p_title2 = doc.add_paragraph()
p_title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p_title2.add_run("Assignment 2 — Multi-Agent System (MAS)")
set_font(run, size=18, bold=True, color=(31, 73, 125))

doc.add_paragraph()

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p_sub.add_run("Healthcare Automated Patient-Care Pipeline")
set_font(run, size=13, italic=True, color=(89, 89, 89))

doc.add_paragraph()

add_table(
    ["Student ID", "Name", "Agent", "Tool"],
    [
        ["IT22248244", "Pandithasundara N B", "PatientIntakeAgent",    "tool_patient_reader"],
        ["IT22014290", "Samishka H T",         "SymptomAnalyzerAgent",  "tool_symptom_analyzer"],
        ["IT22333148", "Wijerathne C G T N",   "TreatmentPlannerAgent", "tool_medication_recommender"],
        ["All members", "All members",         "MedicalReportAgent",    "tool_report_generator"],
    ],
    col_widths=[3, 4.5, 5, 5],
)

doc.add_paragraph()
p_date = doc.add_paragraph()
p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p_date.add_run("Date: April 2026")
set_font(run, size=11, italic=True, color=(89, 89, 89))

doc.add_page_break()


# ============================================================================
# 1. INTRODUCTION
# ============================================================================

heading("1. Introduction")
separator()
para(
    "This report documents the design, implementation, and evaluation of a "
    "Healthcare Multi-Agent System (MAS) developed for SE4010 CTSE Assignment 2. "
    "The system automates the full patient-care pipeline — from record intake to "
    "report generation — using four specialised AI agents that communicate through "
    "a shared state object. The entire system runs locally with no paid cloud API keys.",
    size=11,
)

heading("1.1 Problem Domain", level=2)
para(
    "Managing patient records, matching symptoms to medical conditions, and "
    "producing treatment plans are time-intensive tasks prone to human error. "
    "Healthcare staff must cross-reference symptoms against diagnostic criteria, "
    "screen medications against patient allergies, and document findings in a "
    "structured format. This MAS automates that complete clinical workflow, providing "
    "consistent, reproducible assessments for any patient record in JSON format.",
    size=11,
)

heading("1.2 Objectives", level=2)
bullet("Demonstrate the use of a sequential multi-agent pipeline in a realistic domain.")
bullet("Implement four distinct agents, each owning a dedicated tool.")
bullet("Maintain a shared GlobalState object that accumulates context across agents.")
bullet("Provide full LLMOps observability: every agent action and tool call is logged.")
bullet("Achieve 100% pass rate on 98 unit + integration tests.")

doc.add_page_break()


# ============================================================================
# 2. SYSTEM ARCHITECTURE
# ============================================================================

heading("2. System Architecture")
separator()
para(
    "The system follows a Sequential Coordinator-Worker pipeline pattern. Each agent "
    "runs in a fixed order, reads results left by the previous agent from GlobalState, "
    "performs its specialised task, writes its results back, and passes control to the "
    "next agent. No data is lost between handoffs — the entire pipeline state is "
    "accumulated in one dataclass.",
    size=11,
)

heading("2.1 Pipeline Overview", level=2)

add_table(
    ["Stage", "Agent", "Tool", "Responsibility", "Owner"],
    [
        ["1", "PatientIntakeAgent",    "tool_patient_reader",       "Validate & load patient JSON", "IT22248244"],
        ["2", "SymptomAnalyzerAgent",  "tool_symptom_analyzer",     "Match symptoms → conditions",  "IT22014290"],
        ["3", "TreatmentPlannerAgent", "tool_medication_recommender","Allergy-screened treatment plan","IT22333148"],
        ["4", "MedicalReportAgent",    "tool_report_generator",     "Generate final MD report",     "All members"],
    ],
    col_widths=[1.5, 4.5, 5, 5, 3],
)

heading("2.2 Architecture Diagram", level=2)
screenshot_placeholder(
    "Terminal output showing the pipeline running — paste image here",
    hint="Run: cd ctse_mas && python main.py --patient data/patients/patient_PT001.json  "
         "then press Cmd+Shift+4 (Mac) or Win+Shift+S (Windows) to screenshot the terminal.",
)

heading("2.3 GlobalState Dataclass", level=2)
para(
    "The central communication mechanism is the GlobalState dataclass defined in "
    "config/state.py. Each agent receives the state object, reads from fields populated "
    "by earlier agents, writes its own fields, and returns the updated state.",
    size=11,
)

add_table(
    ["Field Group", "Fields", "Populated By"],
    [
        ["Pipeline Input",       "patient_file_path",                      "main.py"],
        ["Patient Intake",       "patient_info, validation_report, is_valid, validation_errors", "Agent 1"],
        ["Symptom Analysis",     "possible_conditions, risk_level, emergency_flags, top_diagnosis", "Agent 2"],
        ["Treatment Planning",   "treatment_plan, recommended_medications, contraindicated_medications, lifestyle_recommendations", "Agent 3"],
        ["Report Generation",    "final_report, report_path, executive_summary", "Agent 4"],
        ["LLMOps / Observability","agent_logs, pipeline_start, pipeline_end, totals", "All agents"],
    ],
    col_widths=[4, 8, 3],
)

heading("2.4 Orchestration Pattern", level=2)
para(
    "The orchestrator in main.py follows this execution sequence:",
    size=11,
)
code_block("state = reset_state()\nstate = PatientIntakeAgent().run(state)\nstate = SymptomAnalyzerAgent().run(state)\nstate = TreatmentPlannerAgent().run(state)\nstate = MedicalReportAgent().run(state)")
para(
    "If a validation warning is found after Agent 1, the pipeline logs a warning but "
    "continues processing rather than aborting, ensuring partial data still flows to "
    "downstream agents.",
    size=11,
)

doc.add_page_break()


# ============================================================================
# 3. PROJECT STRUCTURE
# ============================================================================

heading("3. Project Structure")
separator()
code_block(
    "ctse_mas/\n"
    "├── main.py                          # Pipeline entry point\n"
    "├── requirements.txt\n"
    "├── conftest.py\n"
    "│\n"
    "├── config/\n"
    "│   ├── state.py                     # GlobalState dataclass\n"
    "│   └── observability.py             # LLMOps logging & Rich console\n"
    "│\n"
    "├── agents/\n"
    "│   ├── agent_patient_intake.py      # Agent 1 — IT22248244\n"
    "│   ├── agent_symptom_analyzer.py    # Agent 2 — IT22014290\n"
    "│   ├── agent_treatment_planner.py   # Agent 3 — IT22333148\n"
    "│   └── agent_report_generator.py   # Agent 4\n"
    "│\n"
    "├── tools/\n"
    "│   ├── tool_patient_reader.py       # Tool 1\n"
    "│   ├── tool_symptom_analyzer.py     # Tool 2\n"
    "│   ├── tool_medication_recommender.py # Tool 3\n"
    "│   └── tool_report_generator.py    # Tool 4\n"
    "│\n"
    "├── data/\n"
    "│   ├── symptoms_db.json             # 10 medical conditions\n"
    "│   ├── medications_db.json          # 14 medications\n"
    "│   └── patients/\n"
    "│       ├── patient_PT001.json       # Influenza case\n"
    "│       └── patient_PT002.json       # UTI case\n"
    "│\n"
    "├── tests/\n"
    "│   ├── test_agent1_patient_intake.py  (21 tests)\n"
    "│   ├── test_agent2_symptom_analyzer.py (22 tests)\n"
    "│   ├── test_agent3_treatment_planner.py (21 tests)\n"
    "│   ├── test_agent4_report_generator.py  (15 tests)\n"
    "│   └── test_pipeline_integration.py    (19 tests)\n"
    "│\n"
    "├── reports/                         # Generated Markdown reports\n"
    "└── logs/                            # LLMOps JSON traces\n"
)

doc.add_page_break()


# ============================================================================
# 4. AGENT DESCRIPTIONS
# ============================================================================

heading("4. Agent Descriptions")
separator()

# ── Agent 1 ──────────────────────────────────────────────────────────────────
heading("4.1 Agent 1 — PatientIntakeAgent (IT22248244)", level=2)
para("Persona: A meticulous clinical administrator responsible for ensuring every patient record is complete and valid before clinical work begins.", italic=True, color=(89, 89, 89))
para("Responsibilities:", bold=True)
bullet("Loads and parses the patient JSON file via tool_patient_reader.")
bullet("Validates required fields: name, age, gender, blood type, vital signs, symptoms.")
bullet("Populates patient_info, validation_report, is_valid, and validation_errors in GlobalState.")
bullet("Produces a structured validation report detailing every field check outcome.")

para("Tool — tool_patient_reader:", bold=True)
para("Reads the JSON file, checks required fields, validates age range (0–120), blood type (8 valid types), vital-sign plausibility (temperature, heart rate, blood pressure), and symptom list. Returns a structured dict with success/failure per field.", size=11)

# ── Agent 2 ──────────────────────────────────────────────────────────────────
heading("4.2 Agent 2 — SymptomAnalyzerAgent (IT22014290)", level=2)
para("Persona: A board-certified clinical diagnostician specialising in evidence-based differential diagnosis. Every conclusion is supported by symptom-to-condition matching evidence.", italic=True, color=(89, 89, 89))
para("Responsibilities:", bold=True)
bullet("Reads symptoms from state.patient_info (populated by Agent 1).")
bullet("Calls tool_symptom_analyzer against symptoms_db.json (10 medical conditions).")
bullet("Populates possible_conditions (ranked list), risk_level, emergency_flags, and top_diagnosis.")
bullet("Immediately escalates any patient with emergency indicators to 'critical' risk.")

para("Tool — tool_symptom_analyzer Algorithm:", bold=True)
para("Computes a normalised overlap score for each condition:", size=11)
code_block("score = matched_symptom_count / max(total_patient_symptoms, total_condition_symptoms)")
para("Conditions with ≥ 1 matched symptom are ranked by confidence %. Emergency indicators are checked independently — any match triggers a risk escalation regardless of confidence score.", size=11)

para("Supported Conditions (symptoms_db.json):", bold=True)
add_table(
    ["Condition", "ICD-10", "Severity", "Urgent Care"],
    [
        ["Influenza",            "J11.1", "Moderate", "No"],
        ["Pneumonia",            "J18.9", "Severe",   "Yes"],
        ["COVID-19",             "U07.1", "Moderate", "Yes"],
        ["Urinary Tract Infection","N39.0","Mild",    "No"],
        ["Diabetes Type 2",      "E11.9", "Moderate", "Yes"],
        ["Hypertension",         "I10",   "Moderate", "Yes"],
        ["Anxiety Disorder",     "F41.1", "Moderate", "No"],
        ["Common Cold",          "J00",   "Mild",     "No"],
        ["Gastroenteritis",      "K59.1", "Mild",     "No"],
        ["Asthma",               "J45.9", "Moderate", "Yes"],
    ],
    col_widths=[5, 2.5, 3, 2.5],
)

# ── Agent 3 ──────────────────────────────────────────────────────────────────
heading("4.3 Agent 3 — TreatmentPlannerAgent (IT22333148)", level=2)
para("Persona: A senior clinical pharmacist who produces safe, evidence-based treatment plans. Never recommends medications that conflict with a patient's known allergies.", italic=True, color=(89, 89, 89))
para("Responsibilities:", bold=True)
bullet("Reads possible_conditions and patient_info (allergies) from GlobalState.")
bullet("Calls tool_medication_recommender against medications_db.json (14 medications).")
bullet("Screens every medication against the patient's allergy list before recommending.")
bullet("Populates treatment_plan, recommended_medications, contraindicated_medications, and lifestyle_recommendations.")

para("Tool — tool_medication_recommender:", bold=True)
para("Matches condition to relevant medications, cross-checks patient allergy fields, excludes any contraindicated drug, and returns a safe medication plan with dosage, route, and side-effect notes.", size=11)

# ── Agent 4 ──────────────────────────────────────────────────────────────────
heading("4.4 Agent 4 — MedicalReportAgent", level=2)
para("Responsibilities:", bold=True)
bullet("Aggregates all pipeline results from GlobalState.")
bullet("Generates a structured Markdown report (saved to reports/) and a JSON summary (saved to reports/).")
bullet("Records a full LLMOps execution trace to logs/.")
bullet("Populates final_report, report_path, and executive_summary in GlobalState.")

doc.add_page_break()


# ============================================================================
# 5. LLMOPS & OBSERVABILITY
# ============================================================================

heading("5. LLMOps & Observability")
separator()
para(
    "Every agent action, tool call, and outcome is captured by the observability module "
    "(config/observability.py). This satisfies the AgentOps & Observability requirement.",
    size=11,
)

heading("5.1 Console Output (Rich)", level=2)
para("Each event is printed to the terminal using the Rich library:", size=11)
add_table(
    ["Event",        "Visual Style",                  "Information Shown"],
    [
        ["Agent Started",   "Green bordered panel",   "Agent name, task, input preview"],
        ["Tool Call",       "Yellow inline log line", "Tool name, input snippet, output snippet"],
        ["Agent Completed", "Blue bordered panel",    "Agent name, duration (seconds), output preview"],
        ["Error",           "Red inline message",     "Agent name, error description"],
    ],
    col_widths=[3.5, 4.5, 6],
)

heading("5.2 JSON Trace Files", level=2)
para(
    "After each pipeline run a timestamped JSON file is written to logs/. "
    "It records every event in order: agent_start, tool_call, and agent_end entries "
    "with full input/output previews and ISO-8601 timestamps.",
    size=11,
)
screenshot_placeholder(
    "Terminal showing full pipeline run with Rich console output",
    hint="Run the pipeline command, scroll to the top of the output, "
         "and take a screenshot of the complete terminal window.",
)

heading("5.3 AgentLog Entries", level=2)
para("Each agent writes a structured AgentLog entry to state.agent_logs:", size=11)
add_table(
    ["Field",          "Description"],
    [
        ["agent_name",    "Identifier of the calling agent"],
        ["action",        "Short description of the task"],
        ["input_summary", "Brief summary of what the agent received"],
        ["output_summary","Brief summary of what the agent produced"],
        ["tool_calls",    "Ordered list of tool names invoked"],
        ["timestamp",     "ISO-8601 timestamp of completion"],
        ["status",        "'success' or 'error'"],
    ],
    col_widths=[4, 10],
)

doc.add_page_break()


# ============================================================================
# 6. SETUP & RUNNING
# ============================================================================

heading("6. Setup & Running the System")
separator()

heading("6.1 Prerequisites", level=2)
bullet("Python 3.10 or later")
bullet("No paid API keys — runs 100% locally")
bullet("Recommended: macOS / Linux terminal or Windows PowerShell")

heading("6.2 Installation", level=2)
code_block(
    "# From the project root:\n"
    "python3 -m venv .venv\n"
    "source .venv/bin/activate          # Windows: .venv\\Scripts\\activate\n"
    "pip install -r ctse_mas/requirements.txt"
)

heading("6.3 Run the Pipeline", level=2)
code_block(
    "cd ctse_mas\n\n"
    "# Influenza patient (PT001)\n"
    "python main.py --patient data/patients/patient_PT001.json\n\n"
    "# UTI patient with comorbidities (PT002)\n"
    "python main.py --patient data/patients/patient_PT002.json"
)
screenshot_placeholder(
    "Pipeline startup banner (green Rich panel with patient file & pipeline stages)",
    hint="Run the pipeline command above and screenshot the first few lines of terminal output.",
)

heading("6.4 Pipeline Output Files", level=2)
add_table(
    ["File", "Location", "Description"],
    [
        ["report_PT001_*.md",   "ctse_mas/reports/", "Full Markdown medical report for PT001"],
        ["report_PT001_*.json", "ctse_mas/reports/", "JSON-structured report summary"],
        ["trace_*.json",        "ctse_mas/logs/",    "Complete LLMOps event trace"],
    ],
    col_widths=[5, 4, 6],
)

screenshot_placeholder(
    "Generated report file opened in VS Code / text editor",
    hint="Open ctse_mas/reports/PT001_medical_report_<latest>.md in VS Code (Markdown preview) and screenshot.",
)

doc.add_page_break()


# ============================================================================
# 7. TESTING
# ============================================================================

heading("7. Testing & Evaluation")
separator()

heading("7.1 Test Suite Overview", level=2)
add_table(
    ["Test File",                          "Owner",     "Tests", "Scope"],
    [
        ["test_agent1_patient_intake.py",    "IT22248244", "21",   "Patient validation, field checks, edge cases"],
        ["test_agent2_symptom_analyzer.py",  "IT22014290", "22",   "Symptom matching, scoring, emergency flags"],
        ["test_agent3_treatment_planner.py", "IT22333148", "21",   "Medication selection, allergy screening"],
        ["test_agent4_report_generator.py",  "All members", "15",  "Report structure, Markdown output"],
        ["test_pipeline_integration.py",     "Group",      "19",   "End-to-end PT001 + PT002 pipeline runs"],
        ["TOTAL",                            "—",          "98",   ""],
    ],
    col_widths=[6, 3, 2, 5],
)

heading("7.2 Running Tests", level=2)
code_block(
    "cd ctse_mas\n\n"
    "# Full suite\n"
    "python -m pytest tests/ -v\n\n"
    "# Individual agent\n"
    "python -m pytest tests/test_agent2_symptom_analyzer.py -v\n\n"
    "# Integration only\n"
    "python -m pytest tests/test_pipeline_integration.py -v"
)
screenshot_placeholder(
    "pytest output showing 98 passed tests",
    hint="Run: cd ctse_mas && python -m pytest tests/ -v  "
         "then screenshot the terminal showing the final '98 passed' summary line.",
)

heading("7.3 Key Test Cases", level=2)
para("Agent 2 (Symptom Analyzer) — Selected Test Scenarios:", bold=True)
add_table(
    ["Test",                          "Description",                          "Expected Result"],
    [
        ["test_influenza_symptoms",     "8 classic flu symptoms",               "Influenza ranked #1, ≥50% confidence"],
        ["test_emergency_flags",        "Symptoms include 'high fever'",        "risk_level = 'critical'"],
        ["test_no_symptoms",            "Empty symptom list",                   "success=False, error returned"],
        ["test_unknown_symptoms",       "Symptoms not in knowledge base",       "Empty conditions list"],
        ["test_uti_patient",            "PT002 UTI patient symptoms",           "UTI ranked #1"],
        ["test_confidence_calculation", "Known symptom set vs. condition size", "score = matched/max(patient, condition)"],
    ],
    col_widths=[5, 5.5, 5.5],
)

doc.add_page_break()


# ============================================================================
# 8. SAMPLE OUTPUT
# ============================================================================

heading("8. Sample Pipeline Output")
separator()
para(
    "The following shows an excerpt from the generated medical report for Patient PT001 "
    "(Amara Perera — Influenza case).",
    size=11,
)

heading("8.1 Executive Summary", level=2)
para(
    "Patient Amara Perera (ID: PT001) was assessed by the Healthcare Multi-Agent System. "
    "The clinical symptom analysis identified 'Influenza' as the most probable diagnosis "
    "(confidence: 57.1%). The overall patient risk level has been classified as CRITICAL. "
    "A total of 3 medication(s) have been recommended after allergy and interaction screening. "
    "URGENT ATTENTION IS REQUIRED — one or more emergency indicators were detected.",
    size=11,
    italic=True,
)

heading("8.2 Differential Diagnosis — PT001", level=2)
add_table(
    ["Rank", "Condition",      "ICD-10", "Confidence", "Severity", "Matched Symptoms"],
    [
        ["1", "Influenza",          "J11.1", "57.1%", "Moderate", "high fever, body aches, fatigue, chills, dry cough, headache, anorexia, sweating"],
        ["2", "Pneumonia [URGENT]", "J18.9", "46.2%", "Severe",   "high fever, fatigue, chills, dry cough, anorexia, sweating"],
        ["3", "Common Cold",        "J00",   "18.2%", "Mild",     "fatigue, dry cough"],
        ["4", "Gastroenteritis",    "K59.1", "18.2%", "Mild",     "fatigue, anorexia"],
        ["5", "Anxiety Disorder",   "F41.1", "16.7%", "Moderate", "fatigue, sweating"],
    ],
    col_widths=[1.5, 4, 2, 2.5, 2.5, 5.5],
)

heading("8.3 Recommended Medications — PT001", level=2)
add_table(
    ["Medication",             "Category",              "Dosage",                     "Notes"],
    [
        ["Paracetamol",          "Analgesic / Antipyretic","500–1000 mg q4–6h PRN",     "First-line; safe for most"],
        ["Ibuprofen",            "NSAID",                 "400–600 mg q6–8h with food", "Take with food"],
        ["Oseltamivir",          "Antiviral",             "75 mg BD × 5 days",          "Best within 48h of onset"],
        ["~~Amoxicillin~~",      "EXCLUDED",              "—",                          "Penicillin allergy"],
    ],
    col_widths=[4, 4, 4.5, 5.5],
)

screenshot_placeholder(
    "Generated Markdown report rendered in VS Code — shows patient demographics and diagnosis table",
    hint="Open the latest .md file in ctse_mas/reports/ in VS Code, "
         "press Cmd+Shift+V (Mac) to open Markdown Preview, then screenshot.",
)

doc.add_page_break()


# ============================================================================
# 9. TECHNICAL REQUIREMENTS CHECKLIST
# ============================================================================

heading("9. Technical Requirements Compliance")
separator()
add_table(
    ["Requirement",                  "Status", "Implementation Detail"],
    [
        ["4 Distinct Agents",          "PASS",   "PatientIntakeAgent, SymptomAnalyzerAgent, TreatmentPlannerAgent, MedicalReportAgent"],
        ["Custom Python Tools",        "PASS",   "4 tools with type hints, docstrings, and structured return dicts"],
        ["State Management",           "PASS",   "GlobalState dataclass passed by reference through all agents"],
        ["Sequential Pipeline",        "PASS",   "Agents run in fixed order; each reads previous agent's output"],
        ["LLMOps / Observability",     "PASS",   "observability.py logs every agent start/end/tool call + JSON trace"],
        ["No Paid APIs",               "PASS",   "Runs entirely locally — no OpenAI/Anthropic/cloud keys needed"],
        ["Individual Agent + Tool",    "PASS",   "Each of the 3 members owns one agent and one corresponding tool"],
        ["Testing & Evaluation",       "PASS",   "98 tests: 21 + 22 + 21 + 15 unit + 19 integration = 98 total"],
        ["Knowledge Base Integration", "PASS",   "symptoms_db.json (10 conditions) + medications_db.json (14 medications)"],
        ["Allergy Screening",          "PASS",   "Agent 3 screens all medications against patient allergy fields"],
    ],
    col_widths=[5, 2, 9],
)

doc.add_page_break()


# ============================================================================
# 10. CONCLUSION
# ============================================================================

heading("10. Conclusion")
separator()
para(
    "This assignment successfully demonstrates the design and implementation of a "
    "healthcare-domain Multi-Agent System. The four-agent sequential pipeline automates "
    "the complete patient-care workflow: record validation, symptom-based differential "
    "diagnosis, allergy-screened treatment planning, and structured report generation.",
    size=11,
)
para(
    "Key technical contributions include:",
    size=11,
)
bullet("A reusable GlobalState dataclass that serves as the single source of truth throughout the pipeline.")
bullet("A normalised overlap scoring algorithm in the SymptomAnalyzerAgent that correctly ranks conditions even when patient symptoms partially overlap multiple conditions.")
bullet("A robust allergy-screening mechanism in TreatmentPlannerAgent that ensures patient safety by excluding contraindicated medications before recommendations are made.")
bullet("A full LLMOps observability layer that records every event in machine-readable JSON format and displays it in human-readable Rich console output simultaneously.")
bullet("A 98-test suite covering unit behaviour, edge cases, and end-to-end pipeline integration.")
para(
    "The system is fully self-contained and locally hosted — no paid APIs or internet "
    "connection are required — making it suitable for secure, air-gapped deployments.",
    size=11,
)

doc.add_page_break()


# ============================================================================
# APPENDIX — SCREENSHOTS GUIDE
# ============================================================================

heading("Appendix A — How to Capture Screenshots for This Report")
separator()
para(
    "The grey placeholder boxes throughout this report mark where screenshots should be inserted. "
    "Follow the steps below to capture and insert each one.",
    size=11,
)

heading("A.1 Running the Pipeline (Terminal Screenshot)", level=2)
code_block(
    "cd /path/to/ctse_mas\n"
    "source ../.venv/bin/activate\n"
    "python main.py --patient data/patients/patient_PT001.json"
)
bullet("Mac: Press Cmd + Shift + 4, drag a rectangle over the terminal window.")
bullet("Windows: Press Win + Shift + S, drag a rectangle, save to clipboard, paste into Paint/Photos.")
bullet("VS Code integrated terminal: right-click → 'Select All' → scroll to see the full output.")
para("Suggested screenshots:", bold=True)
bullet("1. The green startup banner (shows patient file and pipeline order).")
bullet("2. Each agent's green 'Agent Started' and blue 'Agent Completed' Rich panels.")
bullet("3. The final 'Pipeline Complete' summary (if shown).")

heading("A.2 Test Results Screenshot", level=2)
code_block(
    "cd ctse_mas\n"
    "python -m pytest tests/ -v 2>&1 | tail -20"
)
bullet("Run the command above, screenshot the final lines showing '98 passed'.")
bullet("Alternatively screenshot the full pytest -v output scrolled to the bottom.")

heading("A.3 Generated Report Screenshot", level=2)
bullet("Open ctse_mas/reports/ in VS Code File Explorer.")
bullet("Click the latest PT001_medical_report_*.md file.")
bullet("Press Cmd+Shift+V (Mac) or Ctrl+Shift+V (Windows) to open Markdown Preview.")
bullet("Screenshot the preview showing patient demographics, diagnosis table, and treatment plan.")

heading("A.4 Log File Screenshot", level=2)
bullet("Open ctse_mas/logs/trace_<latest>.json in VS Code.")
bullet("Screenshot the first ~30 lines showing the agent_start and tool_call entries.")

heading("A.5 Inserting Screenshots into This Word Document", level=2)
bullet("Replace each grey placeholder box by clicking on it, then Insert → Pictures → select your screenshot file.")
bullet("Right-click the inserted image → 'Wrap Text' → 'In Line with Text' for clean layout.")
bullet("Resize to fit within the page margins (approx. 14 cm wide).")

# ── Save ─────────────────────────────────────────────────────────────────────

output_path = str(Path(__file__).resolve().parent / "CTSE_Assignment2_Report.docx")
doc.save(output_path)
print(f"Report saved to: {output_path}")
